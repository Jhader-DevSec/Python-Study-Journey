from pdf2docx import Converter
from docx import Document
from pathlib import Path
from pypdf import PdfReader

def pdf_converter(input_pdf, output_docx):
    pdf_file = input_pdf
    docx_file = output_docx
    
    cv = Converter(pdf_file)
    
    cv.convert(docx_file, pages='all')
    
    cv.close()
    
def modify_text(file):
    path_extension = Path(file)
    
    if path_extension.suffix.lower() == ".pdf":
        new_name = f"{path_extension.stem}_modificada.docx"
        
        pdf_converter(file, new_name)
        
        document = Document(new_name)
        
        paragraph_number = int(input("qual o paragrafo que você quer modificar? "))
        paragraph = document.paragraphs[paragraph_number]
        paragraph.text = input("digite aqui todo o paragrafo corrigido: ")
        
        document.save(new_name)
        print(f"Sucesso! PDF convertido e salvo como: {new_name}")
        
    elif path_extension.suffix.lower() == ".docx":
        document = Document(file)
        
        paragraph_number = int(input("qual o paragrafo que você quer modificar? "))
        paragraph = document.paragraphs[paragraph_number]
        paragraph.text = input("digite aqui todo o paragrafo corrigido: ")
        
        document.save(file)
        print(f"Sucesso! Arquivo {file} foi modificado.")
        
def read_pdf(file):
    path_extension = Path(file)
    
    if path_extension.suffix.lower() == ".pdf":
        reader = PdfReader(file)
        
        for page in reader.pages:
            yield page.extract_text()
            
def read_docx(file):
    path_extension = Path(file)
    
    if path_extension.suffix.lower() == ".docx":
        reader = Document(file)
        
        for paragraph in reader.paragraphs:
            yield paragraph.text
    
    
while True:
    print("\n----BEM VINDO AO LEITOR E MODIFICADOR DE ARQUIVOS----")
    print("o que deseja fazer?")
    print("1. Ler um documento")
    print("2. Modificar um documento")
    print("3. Converter um documento pdf para docx")
    print("4. sair")
        
    try:
        user_choice = int(input("escolha o número da opção desejada"))
    except ValueError:
        print("por favor digite apenas numeros válidos!")
        continue
    
    if user_choice == 1:
        file_path = input('qual o ducumento desejado para ler?')
        if Path(file_path).suffix.lower() == ".pdf":
            try:
                for page_content in read_pdf(file_path):
                    print(page_content)
                    print("-" * 50)
                    
            except FileNotFoundError:
                print("arquivo não encontrado!")
                
        elif Path(file_path).suffix.lower() == ".docx":
            try:
                for paragraph_content in read_docx(file_path):
                    print(paragraph_content)
                    
            except FileNotFoundError:
                print("arquivo não encontrado!")       
        else:
            print("formato de documento não aceito. Por favor ultilize somente .pdf ou .docx!")
        
    elif user_choice == 2:
        file_path = input('qual o ducumento desejado para modificar?')
        
        try:
            modify_text(file_path)
        except FileNotFoundError:
            print("arquivo não encontrado!")
        except (ValueError, IndexError):
            print("numero de paragrafo invalido!")
        
    elif user_choice == 3:
        file_path = input('qual o documento desejado para converter? ')
        nome_saida = f"{Path(file_path).stem}_convertido.docx"
        
        try:
            pdf_converter(file_path, nome_saida)
            print(f"Convertido com sucesso para {nome_saida}")
        except FileNotFoundError:
            print("arquivo não encontrado!")
        
    elif user_choice == 4:
        print("Obrigado por usar meu codigo!")
        break
    
    else:
        print("por favor selecione o número do 1 ao 4!")
